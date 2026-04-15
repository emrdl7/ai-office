# 산출물 내보내기 엔진 — MD→PDF, MD→DOCX, 사이트→ZIP
from __future__ import annotations
import zipfile
from pathlib import Path


def md_to_html(md_content: str, title: str = '산출물') -> str:
  '''마크다운을 스타일 적용된 HTML로 변환한다.'''
  try:
    import markdown  # type: ignore[import-untyped]
    body = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'nl2br', 'toc'])
  except ImportError:
    body = f'<pre>{md_content}</pre>'

  return f'''<!DOCTYPE html>
<html lang="ko"><head>
<meta charset="utf-8">
<title>{title}</title>
<style>
  body {{ font-family: -apple-system, 'Pretendard', sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.8; color: #333; }}
  h1 {{ font-size: 1.8em; border-bottom: 2px solid #2196F3; padding-bottom: 8px; }}
  h2 {{ font-size: 1.4em; color: #1565C0; margin-top: 2em; }}
  h3 {{ font-size: 1.15em; color: #1976D2; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
  th {{ background: #f5f5f5; }}
  pre {{ background: #f8f9fa; padding: 16px; border-radius: 8px; overflow-x: auto; }}
  code {{ background: #f0f0f0; padding: 2px 6px; border-radius: 4px; font-size: 0.9em; }}
  blockquote {{ border-left: 4px solid #2196F3; margin: 1em 0; padding: 0.5em 1em; background: #f8f9fa; }}
</style>
</head><body>{body}</body></html>'''


def md_to_pdf(md_content: str, pdf_path: str | Path, title: str = '산출물') -> Path:
  '''마크다운 → HTML → PDF 변환.'''
  pdf_path = Path(pdf_path)
  pdf_path.parent.mkdir(parents=True, exist_ok=True)

  html = md_to_html(md_content, title)

  try:
    import weasyprint
    weasyprint.HTML(string=html).write_pdf(str(pdf_path))
    return pdf_path
  except ImportError:
    # weasyprint 없으면 HTML로 저장
    html_path = pdf_path.with_suffix('.html')
    html_path.write_text(html, encoding='utf-8')
    return html_path


def md_to_docx(md_content: str, docx_path: str | Path, title: str = '산출물') -> Path:
  '''마크다운 → DOCX 변환.'''
  from docx import Document
  from docx.shared import Pt

  docx_path = Path(docx_path)
  docx_path.parent.mkdir(parents=True, exist_ok=True)

  doc = Document()
  doc.add_heading(title, 0)

  for line in md_content.split('\n'):
    stripped = line.strip()
    if stripped.startswith('### '):
      doc.add_heading(stripped[4:], level=3)
    elif stripped.startswith('## '):
      doc.add_heading(stripped[3:], level=2)
    elif stripped.startswith('# '):
      doc.add_heading(stripped[2:], level=1)
    elif stripped.startswith('- ') or stripped.startswith('* '):
      doc.add_paragraph(stripped[2:], style='List Bullet')
    elif stripped:
      # 볼드/이탤릭 마크다운은 그대로 텍스트로
      doc.add_paragraph(stripped)

  doc.save(str(docx_path))
  return docx_path


def folder_to_zip(folder_path: str | Path, zip_path: str | Path) -> Path:
  '''폴더를 ZIP으로 압축한다.'''
  folder_path = Path(folder_path)
  zip_path = Path(zip_path)
  zip_path.parent.mkdir(parents=True, exist_ok=True)

  with zipfile.ZipFile(str(zip_path), 'w', zipfile.ZIP_DEFLATED) as zf:
    for file in folder_path.rglob('*'):
      if file.is_file() and not file.name.startswith('.'):
        zf.write(file, file.relative_to(folder_path))
  return zip_path


def get_exportable_formats(task_dir: Path) -> list[str]:
  '''태스크 디렉토리에서 내보내기 가능한 포맷 목록을 반환한다.'''
  formats: list[str] = []
  if not task_dir.exists():
    return formats

  has_md = any(task_dir.rglob('*.md'))
  has_html = any(task_dir.rglob('*.html'))

  if has_md:
    formats.extend(['pdf', 'docx'])
  if has_html:
    formats.append('pdf')
  if any(task_dir.iterdir()):
    formats.append('zip')

  return list(set(formats))
